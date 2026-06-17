from __future__ import annotations

from datetime import date, datetime, time
from html.parser import HTMLParser
from pathlib import Path
from typing import Any
import csv
import json
import re
import xml.etree.ElementTree as ET

from .model import SourceDocument


class _HTMLText(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        text = data.strip()
        if text:
            self.parts.append(text)

    def text(self) -> str:
        return "\n".join(self.parts)


def _normalize(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip()


def extract_text(path: str | Path) -> SourceDocument:
    p = Path(path).expanduser().resolve()
    if not p.exists() or not p.is_file():
        raise FileNotFoundError(f"Source file not found: {p}")
    suffix = p.suffix.lower()
    metadata: dict[str, Any] = {"extractor": "plain"}

    if suffix in {".txt", ".md", ".log", ".py", ".js", ".ts"}:
        text = p.read_text(encoding="utf-8", errors="replace")
    elif suffix in {".html", ".htm"}:
        raw = p.read_text(encoding="utf-8", errors="replace")
        parser = _HTMLText()
        parser.feed(raw)
        text = parser.text()
        metadata = {"extractor": "html.parser"}
    elif suffix == ".csv":
        text, metadata = _csv_to_text(p.read_text(encoding="utf-8", errors="replace"))
    elif suffix == ".json":
        text, metadata = _json_to_text(p.read_text(encoding="utf-8", errors="replace"))
    elif suffix == ".xml":
        text, metadata = _xml_to_text(p.read_text(encoding="utf-8", errors="replace"))
    elif suffix == ".pdf":
        text, metadata = _pdf_to_text(p)
    elif suffix == ".docx":
        text, metadata = _docx_to_text(p)
    elif suffix in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
        text, metadata = _excel_to_text(p)
    else:
        raw = p.read_bytes()
        text = raw.decode("utf-8", errors="replace")
        metadata = {"extractor": "utf8-fallback", "warning": "Unknown extension; decoded as UTF-8 with replacement."}

    return SourceDocument.from_path(p, _normalize(text), metadata)


def _csv_to_text(raw: str) -> tuple[str, dict[str, Any]]:
    rows = list(csv.reader(raw.splitlines()))
    lines = ["# CSV document"]
    header = rows[0] if rows else []
    if header:
        lines.append("Header: " + " | ".join(header))
    for idx, row in enumerate(rows[1:] if header else rows, start=2 if header else 1):
        if header:
            pairs = []
            for col, value in zip(header, row):
                if value.strip():
                    pairs.append(f"{col.strip()}={value.strip()}")
            lines.append(f"Row {idx}: " + " | ".join(pairs))
        else:
            lines.append(f"Row {idx}: " + " | ".join(row))
    return "\n".join(lines), {"extractor": "csv", "rows": len(rows), "columns": len(header)}


def _json_to_text(raw: str) -> tuple[str, dict[str, Any]]:
    try:
        obj = json.loads(raw)
    except json.JSONDecodeError:
        return raw, {"extractor": "json-raw", "warning": "Invalid JSON; raw text preserved."}
    return json.dumps(obj, ensure_ascii=False, indent=2), {"extractor": "json", "type": type(obj).__name__}


def _xml_to_text(raw: str) -> tuple[str, dict[str, Any]]:
    try:
        root = ET.fromstring(raw)
    except ET.ParseError:
        return raw, {"extractor": "xml-raw", "warning": "Invalid XML; raw text preserved."}
    parts: list[str] = [f"# XML root: {root.tag}"]
    for elem in root.iter():
        if elem.text and elem.text.strip():
            parts.append(f"{elem.tag}: {elem.text.strip()}")
    return "\n".join(parts), {"extractor": "xml", "root": root.tag}


def _pdf_to_text(path: Path) -> tuple[str, dict[str, Any]]:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception as exc:
        raise RuntimeError("PDF support requires pypdf. Run: pip install pypdf") from exc
    reader = PdfReader(str(path))
    pages = []
    for idx, page in enumerate(reader.pages, start=1):
        pages.append(f"--- page {idx} ---\n{page.extract_text() or ''}")
    return "\n\n".join(pages), {"extractor": "pypdf", "pages": len(reader.pages)}


def _docx_to_text(path: Path) -> tuple[str, dict[str, Any]]:
    try:
        import docx  # type: ignore
    except Exception as exc:
        raise RuntimeError("DOCX support requires python-docx. Run: pip install python-docx") from exc
    doc = docx.Document(str(path))
    parts = [f"# DOCX document: {path.name}"]
    parts.extend(p.text for p in doc.paragraphs if p.text.strip())
    for table_index, table in enumerate(doc.tables, start=1):
        parts.append(f"--- table {table_index} ---")
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts), {"extractor": "python-docx", "paragraphs": len(doc.paragraphs), "tables": len(doc.tables)}


def _cell_to_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.isoformat(sep=" ", timespec="seconds")
    if isinstance(value, date):
        return value.isoformat()
    if isinstance(value, time):
        return value.isoformat(timespec="seconds")
    return str(value).strip()


def _trim(values: list[str]) -> list[str]:
    last = -1
    for idx, value in enumerate(values):
        if value:
            last = idx
    return values[: last + 1] if last >= 0 else []


def _excel_to_text(path: Path) -> tuple[str, dict[str, Any]]:
    try:
        from openpyxl import load_workbook  # type: ignore
    except Exception as exc:
        raise RuntimeError("Excel support requires openpyxl. Run: pip install openpyxl") from exc

    wb = load_workbook(filename=str(path), read_only=True, data_only=False)
    parts = [f"# Excel workbook: {path.name}", f"Sheets: {', '.join(wb.sheetnames)}"]
    sheets: list[dict[str, Any]] = []
    total_rows = 0
    total_cells = 0

    for sheet in wb.worksheets:
        parts.append("")
        parts.append(f"## Sheet: {sheet.title}")
        dim = sheet.calculate_dimension(force=True)
        parts.append(f"Dimension: {dim}")
        header: list[str] | None = None
        rows_with_data = 0
        cells_with_data = 0
        for row_idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            values = _trim([_cell_to_text(v) for v in row])
            if not values:
                continue
            rows_with_data += 1
            cells_with_data += len([v for v in values if v])
            if header is None:
                header = [value or f"Column {i}" for i, value in enumerate(values, start=1)]
                parts.append("Header: " + " | ".join(header))
                continue
            pairs: list[str] = []
            for col_idx, value in enumerate(values, start=1):
                if not value:
                    continue
                key = header[col_idx - 1] if col_idx <= len(header) else f"Column {col_idx}"
                pairs.append(f"{key}={value}")
            if pairs:
                parts.append(f"Row {row_idx}: " + " | ".join(pairs))
        total_rows += rows_with_data
        total_cells += cells_with_data
        sheets.append({"name": sheet.title, "dimension": dim, "rows_with_data": rows_with_data, "cells_with_data": cells_with_data})
    wb.close()
    return "\n".join(parts), {
        "extractor": "openpyxl",
        "workbook_type": "excel",
        "sheet_count": len(sheets),
        "sheets": sheets,
        "rows_with_data": total_rows,
        "cells_with_data": total_cells,
    }
